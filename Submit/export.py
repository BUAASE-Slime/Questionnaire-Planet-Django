

from Qn.models import *
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx import *
from io import BytesIO

import djangoProject.settings



def paper_to_docx(qn_id):
    document = Document()
    survey = Survey.objects.get(survey_id=qn_id)
    docx_title = survey.title + '_' + str(survey.username) + '_' + str(qn_id)+"考卷" + ".docx"

    print(docx_title)

    # run = document.add_paragraph().add_run('This is a letter.')
    # font = run.font
    # font.name = '宋体' 英文字体设置
    document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
    document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_paragraph().add_run('第二个段落，abcDEFg，这是中文字符', style='Song')

    document.add_heading(survey.title, 0)


    paragraph = document.add_paragraph().add_run("考试介绍： "+survey.description, style='Song')
    sum_score = 0
    questions = Question.objects.filter(survey_id=survey)
    for question in questions:
        sum_score += question.point
    introduction = "考试须知：本考卷共计" + str(survey.question_num) + "个问题，总分共计 "+str(sum_score)+"分"

    paragraph = document.add_paragraph().add_run(introduction, style='Song')
    str_time = survey.finished_time
    if survey.finished_time  is None:
        str_time = "暂未设定"

    warning = "此外本场考试的截止时间为：" + str(str_time) + "。注意不要在考试截止时间后提交试卷！！"
    document.add_paragraph().add_run(warning, style='Song')
    questions = [] ; question_info_list = [];
    question_list = Question.objects.filter(survey_id=survey)
    for question in question_list:
        if question.type in ['school','name','class','stuId']:
            question_info_list.append(question)
        else:
            questions.append(question)
    for question in question_info_list:
        document.add_paragraph().add_run(question.title, style='Song')
    i = 1
    for question in questions:

        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "单选题"
        elif type == 'checkbox':
            type_str = '多选题'
        elif type == 'text':
            type_str = '填空题'
        elif type == 'mark':
            type_str = '评分题'
        elif type == 'judge':
            type_str = "判断题"
        document.add_paragraph().add_run(str(i) + "、" + question.title + "(" + type_str +"  "+ str(question.point)+"分 )", style='Song')

        i += 1
        options = Option.objects.filter(question_id=question)
        option_option = 0
        num = 1
        for option in options:
            option_str = "      "

            alphas = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if question.type in ['checkbox', 'radio']:
                option_str += alphas[option_option] + " :  "
                # option_str += "选项 " + str(num) + " :  "
                option_option += 1
                num += 1

            option_str += option.content
            document.add_paragraph().add_run(option_str, style='Song')
        if question.type in ['mark', 'text']:
            document.add_paragraph(' ')

    document.add_page_break()
    # document.add_paragraph(str(qn_id))
    f = BytesIO()
    save_path = docx_title

    document.save(f)
    # document.save(save_path)

    docx_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        docx_path = djangoProject.settings.MEDIA_ROOT + "/Document/"

    print(docx_path)
    document.save(docx_path + docx_title)

    return document, f, docx_title, docx_path

import xlwt

from Qn.views import KEY_STR
def write_exam_to_excel(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    submit_list = Submit.objects.filter(survey_id=qn)

    xls = xlwt.Workbook()
    sht1 = xls.add_sheet("Sheet1")

    sht1.write(0, 0, "序号")
    sht1.write(0, 1, "提交者用户名")
    sht1.write(0, 2, "提交时间")
    question_list = Question.objects.filter(survey_id=qn)
    question_info_list = []; questions = []
    paper_sum_score = 0
    for question in question_list:
        paper_sum_score += question.point
        if question.type in ['school','name','class','stuId']:
            question_info_list.append(question)
        else:# TODO 都是问题吧？
            questions.append(question)
    i = 1
    for question in question_info_list:
        sht1.write(0, 2 + i,  question.title)
        i += 1

    question_num = len(questions)
    info_num = len(question_info_list)
    # i = 1

    for question in questions:
        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "单选题"
        elif type == 'checkbox':
            type_str = '多选题'
        elif type == 'text':
            type_str = '填空题'
        elif type == 'mark':
            type_str = '评分题'
        elif type == 'judge':
            type_str = "判断题"
        sht1.write(0, 2 + i, str(i-info_num) + "、" + question.title+" ("+type_str+" "+str(question.point)+"分)")
        i += 1
    sht1.write(0, 2 + i, "总分 ("+str(paper_sum_score)+"分)")
    pattern_green = xlwt.Pattern()  # Create the Pattern
    pattern_green.pattern = xlwt.Pattern.SOLID_PATTERN  # May be: NO_PATTERN, SOLID_PATTERN, or 0x00 through 0x12
    pattern_green.pattern_fore_colour = 3
    # May be: 8 through 63. 0 = Black, 1 = White, 2 = Red, 3 = Green, 4 = Blue, 5 = Yellow, 6 = Magenta, 7 = Cyan, 16 = Maroon, 17 = Dark Green, 18 = Dark Blue, 19 = Dark Yellow , almost brown), 20 = Dark Magenta, 21 = Teal, 22 = Light Gray, 23 = Dark Gray, the list goes on...
    style_green = xlwt.XFStyle()  # Create the Pattern
    # style_green.pattern = pattern_green # Add Pattern to Style
    font_green = xlwt.Font()
    font_green.colour_index = 17
    style_green.font = font_green
    pattern_red = xlwt.Pattern()  # Create the Pattern
    pattern_red.pattern = xlwt.Pattern.SOLID_PATTERN  # May be: NO_PATTERN, SOLID_PATTERN, or 0x00 through 0x12
    pattern_red.pattern_fore_colour = 2
    # May be: 8 through 63. 0 = Black, 1 = White, 2 = Red, 3 = Green, 4 = Blue, 5 = Yellow, 6 = Magenta, 7 = Cyan, 16 = Maroon, 17 = Dark Green, 18 = Dark Blue, 19 = Dark Yellow , almost brown), 20 = Dark Magenta, 21 = Teal, 22 = Light Gray, 23 = Dark Gray, the list goes on...
    style_red = xlwt.XFStyle()  # Create the Pattern
    # style_red.pattern = pattern_red  # Add Pattern to Style
    font_red = xlwt.Font()
    font_red.colour_index = 2
    style_red.font = font_red
    sht1.write(1, 0, "正确答案")
    i = 2+info_num+1
    for question in questions:
        answer_str = question.right_answer
        answer_str = answer_str.replace(KEY_STR, ';')
        sht1.write(1, i, answer_str)
        i+=1

    id = 2
    for submit in submit_list:
        sht1.write(id, 0, id-1)
        username = submit.username
        if username == '' or username is None:
            username = "匿名用户"
        sht1.write(id, 1, username)
        sht1.write(id, 2, submit.submit_time.strftime("%Y/%m/%d %H:%M"))
        question_num = 1
        for question in question_info_list:
            answer_str = (Answer.objects.get(submit_id=submit, question_id=question)).answer
            sht1.write(id, 2 + question_num, answer_str)
            question_num += 1
        personal_score = 0
        for question in questions:

            answer_str = ""
            try:
                answer = Answer.objects.get(submit_id=submit, question_id=question)
                answer_str = answer.answer
            except:
                answer_str = ""
            if question.type == 'checkbox':
                answer_str = answer_str.replace(KEY_STR,';')
            if answer.answer == question.right_answer:
                answer.score = question.point
                style = style_green
            else:
                style = style_red
                answer.score = 0
            answer.save()
            personal_score+=answer.score

            sht1.write(id, 2 + question_num, answer_str,style)

            question_num += 1
        submit.score = personal_score
        submit.save()
        sht1.write(id, 2 + question_num, submit.score)

        id += 1
    save_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        save_path = djangoProject.settings.MEDIA_ROOT + "/Document/"
    excel_name = qn.title + "问卷的统计信息" + ".xls"
    xls.save(save_path + excel_name)
    return excel_name

def vote_to_docx(qn_id):
    document = Document()
    survey = Survey.objects.get(survey_id=qn_id)
    docx_title = survey.title + '_' + str(survey.username) + '_' + str(qn_id) + ".docx"

    # code = hash_code(str(survey.username),str(qn_id))

    # docx_title = code
    print(docx_title)

    # run = document.add_paragraph().add_run('This is a letter.')
    # font = run.font
    # font.name = '宋体' 英文字体设置
    document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
    document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_paragraph().add_run('第二个段落，abcDEFg，这是中文字符', style='Song')

    document.add_heading(survey.title, 0)

    paragraph_list = []

    paragraph = document.add_paragraph().add_run(survey.description, style='Song')

    introduction = "本投票问卷已经收集了" + str(survey.recycling_num) + "份，共计" + str(survey.question_num) + "个问题"
    paragraph = document.add_paragraph().add_run(introduction, style='Song')
    paragraph_list.append(paragraph)

    questions = Question.objects.filter(survey_id=survey)
    i = 1
    for question in questions:
        vote_str = ""
        if question.isVote:
            vote_str = "投票"
        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "单选题"
        elif type == 'checkbox':
            type_str = '多选题'
        elif type == 'text':
            type_str = '填空题'
        elif type == 'mark':
            type_str = '评分题'
        document.add_paragraph().add_run(str(i) + "、" + question.title + "(" + vote_str+type_str + ")", style='Song')

        i += 1
        options = Option.objects.filter(question_id=question)
        option_option = 0
        num = 1
        for option in options:
            option_str = "      "

            alphas = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if question.type in ['checkbox', 'radio']:
                # option_str += alphas[option_option] + " :  "
                option_str += "选项 " + str(num) + " :  "
                option_option += 1
                num += 1

            option_str += option.content
            document.add_paragraph().add_run(option_str, style='Song')
        if question.type in ['mark', 'text']:
            document.add_paragraph(' ')

    document.add_page_break()
    # document.add_paragraph(str(qn_id))
    f = BytesIO()
    save_path = docx_title

    document.save(f)
    # document.save(save_path)

    docx_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        docx_path = djangoProject.settings.MEDIA_ROOT + "/Document/"

    print(docx_path)
    document.save(docx_path + docx_title)

    return document, f, docx_title, docx_path


def write_vote_to_excel(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    submit_list = Submit.objects.filter(survey_id=qn)

    xls = xlwt.Workbook()
    sht1 = xls.add_sheet("Sheet1")

    sht1.write(0, 0, "序号")
    sht1.write(0, 1, "提交者")
    sht1.write(0, 2, "提交时间")
    question_list = Question.objects.filter(survey_id=qn)
    question_sum = len(question_list)
    i = 1

    for question in question_list:
        sht1.write(0, 2 + i, str(i) + "、" + question.title)
        i += 1

    id = 1
    for submit in submit_list:
        sht1.write(id, 0, id)
        username = submit.username
        if username == '' or username is None:
            username = "匿名用户"
        sht1.write(id, 1, username)
        sht1.write(id, 2, submit.submit_time.strftime("%Y/%m/%d %H:%M"))
        question_num = 1
        for question in question_list:
            answer_str = ""
            try:
                answer = Answer.objects.get(submit_id=submit, question_id=question)
                answer_str = answer.answer
            except:
                answer_str = ""
            if question.type == 'checkbox':
                answer_str = answer_str.replace(KEY_STR, ';')

            sht1.write(id, 2 + question_num, answer_str)

            question_num += 1

        id += 1
    question_num = 1
    submit_num = len(submit_list)
    option_id = 0

    for question in question_list:
        if question.isVote:
            result_str = ""
            option_list = Option.objects.filter(question_id=question)
            answer_list = Answer.objects.filter(question_id=question)
            option_max_num = 0
            for option in option_list:
                option_num = 0
                content = option.content
                for answer in answer_list:
                    if answer.answer.find(content) >= 0:
                        option_num += 1
                if option_num > option_max_num:
                    option_max_num = option_num
                    result_str = content
                    option_id = option.option_id
            for option in option_list:
                option_num = 0
                content = option.content
                for answer in answer_list:
                    if answer.answer.find(content) >= 0:
                        option_num += 1
                if option_num == option_max_num and option_id != option.option_id:
                    option_max_num = option_num
                    result_str += ";"+ content
            sht1.write(id,2+question_num,result_str)
            sht1.write(id+1, 2 + question_num, option_max_num)
        question_num += 1


    sht1.write(id, 0, "投票最高结果")
    sht1.write(id+1, 0, "投票最高人数")
    save_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    from .views import IS_LINUX
    if IS_LINUX:
        save_path = djangoProject.settings.MEDIA_ROOT + "/Document/"
    excel_name = qn.title + "问卷的统计信息" + ".xls"
    xls.save(save_path + excel_name)
    return excel_name